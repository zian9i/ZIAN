from docx import Document

# Create a new Document
doc = Document()

# Add Title
doc.add_heading('Resume', 0)

# Add Name and Contact Info
doc.add_paragraph('Name: Zain')
doc.add_paragraph('Date of Birth: Not Provided')
doc.add_paragraph('Place of Residence: Russia')
doc.add_paragraph('Location: Not Specified')
doc.add_paragraph('Phone Number: Not Provided')
doc.add_paragraph('Email: [Email Address]')

# Add Education Section
doc.add_heading('Education', level=1)
doc.add_paragraph('University: [University Name]')
doc.add_paragraph('Major: [Field of Study]')
doc.add_paragraph('GPA: [If Applicable]')
doc.add_paragraph('Years: [Start Year - Graduation Year]')

# Add Skills Section
doc.add_heading('Skills', level=1)
doc.add_paragraph('Languages:')
doc.add_paragraph('  - Arabic (Native)')
doc.add_paragraph('  - English (B2 - Upper Intermediate)')

doc.add_paragraph('Technical Skills:')
doc.add_paragraph('  - Programming Languages: Python, C++, JavaScript')
doc.add_paragraph('  - Operating Systems: Windows, Linux, macOS')
doc.add_paragraph('  - Virtualization: Virtual Machines, experienced with environments like Kali Linux and Kali Purple')
doc.add_paragraph('  - Computer Technology: Strong knowledge of hardware, software, networking, and cybersecurity')

doc.add_paragraph('Additional Skills:')
doc.add_paragraph('  - Device repairs and tool usage (e.g., soldering iron, wire work, screwdrivers)')

# Add Professional Experience Section
doc.add_heading('Professional Experience', level=1)
doc.add_paragraph('YouTube Channel (6+ Years):')
doc.add_paragraph('  Managed a YouTube channel focusing on inventions and creative projects. Took a break from content creation approximately 3 years ago.')

# Add Projects and Interests Section
doc.add_heading('Projects & Interests', level=1)
doc.add_paragraph('  Passionate about electronics and innovation, with a focus on exploring, disassembling, and understanding devices.')
doc.add_paragraph('  Actively involved in technology-driven projects and continuous learning in cybersecurity and programming.')

# Save the document
file_path = "/mnt/data/Zain_Resume.docx"
doc.save(file_path)

file_path
